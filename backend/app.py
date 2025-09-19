from flask import Flask, request, send_file, jsonify
from flask_cors import CORS
from docx import Document
import io
import json
import os
import subprocess
import tempfile
import logging
import shutil

app = Flask(__name__)
CORS(app)

# Set up logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')

@app.route("/", methods=["GET"])
def root():
    return jsonify({
        "status": "success", 
        "message": "Cover Letter Editor API is running"
    })

@app.route("/health", methods=["GET"])
def health_check():
    # Test if LibreOffice is working
    try:
        result = subprocess.run(['libreoffice', '--version'], capture_output=True, text=True, timeout=10)
        libreoffice_status = "working" if result.returncode == 0 else "not working"
        return jsonify({
            "status": "healthy",
            "service": "cover-letter-editor",
            "libreoffice": libreoffice_status,
            "libreoffice_version": result.stdout.strip() if result.returncode == 0 else result.stderr
        })
    except Exception as e:
        return jsonify({
            "status": "unhealthy",
            "error": f"LibreOffice check failed: {str(e)}"
        }), 500

def convert_docx_to_pdf(docx_path, output_dir):
    """
    Convert DOCX to PDF using LibreOffice with better error handling
    """
    try:
        logging.info(f"Starting conversion: {docx_path} -> PDF in {output_dir}")
        
        # Verify the input file exists
        if not os.path.exists(docx_path):
            logging.error(f"Input file does not exist: {docx_path}")
            return None
            
        # Verify output directory exists and is writable
        if not os.path.exists(output_dir):
            os.makedirs(output_dir)
            
        # Use a more direct approach with subprocess.run
        cmd = [
            'libreoffice', 
            '--headless', 
            '--convert-to', 
            'pdf', 
            '--outdir', 
            output_dir, 
            docx_path
        ]
        
        logging.info(f"Running command: {' '.join(cmd)}")
        
        result = subprocess.run(
            cmd,
            capture_output=True,
            text=True,
            timeout=120,  # Longer timeout for conversion
            cwd=output_dir  # Run in the output directory
        )
        
        logging.info(f"LibreOffice return code: {result.returncode}")
        logging.info(f"LibreOffice stdout: {result.stdout}")
        if result.stderr:
            logging.error(f"LibreOffice stderr: {result.stderr}")
        
        if result.returncode == 0:
            # Find the generated PDF file
            expected_pdf_name = os.path.basename(docx_path).replace('.docx', '.pdf')
            expected_pdf_path = os.path.join(output_dir, expected_pdf_name)
            
            if os.path.exists(expected_pdf_path):
                file_size = os.path.getsize(expected_pdf_path)
                logging.info(f"PDF successfully created: {expected_pdf_path} ({file_size} bytes)")
                return expected_pdf_path
            else:
                # Check what files were actually created
                created_files = os.listdir(output_dir)
                logging.error(f"Expected PDF not found. Files in directory: {created_files}")
                
                # Look for any PDF files that might have been created
                pdf_files = [f for f in created_files if f.endswith('.pdf')]
                if pdf_files:
                    logging.info(f"Found PDF files: {pdf_files}")
                    return os.path.join(output_dir, pdf_files[0])
                
                return None
        else:
            logging.error(f"LibreOffice conversion failed with return code {result.returncode}")
            return None
            
    except subprocess.TimeoutExpired:
        logging.error("LibreOffice conversion timed out after 120 seconds")
        return None
    except Exception as e:
        logging.error(f"Unexpected error during conversion: {str(e)}")
        return None

@app.route("/edit", methods=["POST"])
def edit_document():
    temp_dir = None
    try:
        logging.info("Edit request received")
        
        # Get uploaded file
        if 'file' not in request.files:
            return {"error": "No file provided"}, 400
            
        file = request.files['file']
        if file.filename == '' or not file.filename.lower().endswith('.docx'):
            return {"error": "Please upload a valid .docx file"}, 400
            
        logging.info(f"File received: {file.filename}")

        # Get replacements from frontend
        replacements = json.loads(request.form["replacements"])
        pdf_name = request.form.get("pdfName", "EditedCoverLetter")
        logging.info(f"Processing with {len(replacements)} replacements")

        # Load document into memory
        doc = Document(file)
        logging.info("Document loaded successfully")

        # Replace text in paragraphs
        replacement_count = 0
        for p in doc.paragraphs:
            for r in replacements:
                for run in p.runs:
                    if r["find"] in run.text:
                        run.text = run.text.replace(r["find"], r["replace"])
                        replacement_count += 1

        # Replace text inside tables 
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for r in replacements:
                        if r["find"] in cell.text:
                            cell.text = cell.text.replace(r["find"], r["replace"])
                            replacement_count += 1

        logging.info(f"Text replacement completed: {replacement_count} replacements made")

        # Create temporary directory for all files
        temp_dir = tempfile.mkdtemp()
        logging.info(f"Created temp directory: {temp_dir}")
        
        docx_path = os.path.join(temp_dir, "document.docx")
        
        # Save the document
        doc.save(docx_path)
        file_size = os.path.getsize(docx_path)
        logging.info(f"Document saved to: {docx_path} ({file_size} bytes)")

        # Convert using LibreOffice
        logging.info("Starting PDF conversion...")
        converted_pdf_path = convert_docx_to_pdf(docx_path, temp_dir)
        
        if not converted_pdf_path:
            logging.error("PDF conversion failed")
            return {"error": "PDF conversion failed. The document may be corrupted or contain unsupported elements."}, 500

        # Verify the PDF was created and has content
        pdf_size = os.path.getsize(converted_pdf_path)
        if pdf_size == 0:
            logging.error("PDF file is empty (0 bytes)")
            return {"error": "Generated PDF is empty"}, 500
            
        logging.info(f"PDF created successfully: {pdf_size} bytes")

        # Read the converted PDF
        with open(converted_pdf_path, "rb") as f:
            pdf_bytes = io.BytesIO(f.read())

        pdf_bytes.seek(0)

        logging.info("Returning PDF response to client")
        return send_file(
            pdf_bytes,
            as_attachment=True,
            download_name=f"{pdf_name}.pdf",
            mimetype="application/pdf",
        )
    
    except Exception as e:
        logging.error(f"Error in edit_document: {str(e)}", exc_info=True)
        return {"error": f"Internal server error: {str(e)}"}, 500
        
    finally:
        # Clean up temporary files
        if temp_dir and os.path.exists(temp_dir):
            try:
                shutil.rmtree(temp_dir)
                logging.info("Temporary files cleaned up")
            except Exception as e:
                logging.warning(f"Error cleaning up files: {str(e)}")

if __name__ == "__main__":
    app.run(host="0.0.0.0", port=5000, debug=False)